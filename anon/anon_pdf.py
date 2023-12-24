import spacy
import docx

def should_anonymize(token):
    # Additional checks to determine whether to anonymize the token
    return token.ent_type_ and token.ent_type_ not in ["PERSON"] or token.like_num

def anonymize_text(text, nlp):
    doc = nlp(text)
    return ''.join([' XXX ' if should_anonymize(token) else token.text + token.whitespace_ for token in doc])

def process_paragraph(paragraph, nlp):
    paragraph.text = anonymize_text(paragraph.text, nlp)
    for run in paragraph.runs:
        run.text = anonymize_text(run.text, nlp)

def process_cell(cell, nlp):
    cell.text = anonymize_text(cell.text, nlp)
    for paragraph in cell.paragraphs:
        process_paragraph(paragraph, nlp)

def anonymize_document(docx_file):
    nlp = spacy.load("ru_core_news_sm")
    document = docx.Document(docx_file)

    for paragraph in document.paragraphs:
        process_paragraph(paragraph, nlp)

    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                process_cell(cell, nlp)

    document.save(docx_file)

# file_path = r"anon\examples\dist_obrabotka_full.docx"
# anonymize_document(file_path)